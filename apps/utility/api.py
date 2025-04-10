import datetime
import time
import json
import os
import base64
import urllib.parse
import logging
import io

import requests
import wechatpy
from wechatpy.session.redisstorage import RedisStorage
from django_redis import get_redis_connection
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage
from rest_framework.response import Response
from rest_framework.viewsets import GenericViewSet
from rest_framework.permissions import IsAuthenticatedOrReadOnly, AllowAny, IsAuthenticated
from baseconfig.models import BaseConfigFileUpload
from system.models import SMSLog
from . import office
from . import serializers


logger = logging.getLogger('restapi')


class WechatJsSign(GenericViewSet):
    """
    微信JS-SDK签名接口

    用于获取微信JS接口的签名配置信息

    请求参数:
        sys_id (int): 系统ID
        url (str): 需要签名的URL

    返回数据:
        {
            "appId": "公众号appid",
            "nonceStr": "随机字符串",
            "timestamp": "时间戳",
            "signature": "签名"
        }
    """
    queryset = BaseConfigFileUpload.objects.none()
    serializer_class = serializers.WechatJsSign
    permission_classes = [AllowAny]

    def create(self, request, *args, **kwargs):
        from system.models import WechatConfig
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = request.data
        weconf = WechatConfig.objects.get(system__sys_id=data['sys_id'])
        redis_client = get_redis_connection("default")
        session_interface = RedisStorage(
            redis_client,
            prefix="wechatpy",
        )
        cli = wechatpy.WeChatClient(weconf.mp_aid, weconf.mp_sk, session=session_interface)
        ticket = cli.jsapi.get_jsapi_ticket()
        notice = str(time.time())
        timestamp = str(int(time.time()))
        url = data['url']
        sign = cli.jsapi.get_jsapi_signature(notice, ticket, timestamp, url)
        return Response({
            "appId": weconf.mp_aid,
            "nonceStr": notice,
            "timestamp": timestamp,
            "signature": sign
        })


class WechatAppQrCode(GenericViewSet):
    """
    微信小程序码生成接口

    用于生成微信小程序的永久二维码

    请求参数:
        sys_id (int): 系统ID
        scene (str): 场景值参数
        page (str, optional): 跳转页面路径,默认pages/index/index

    返回数据:
        {
            "qrcode": "小程序码图片URL"
        }
    """
    queryset = BaseConfigFileUpload.objects.none()
    serializer_class = serializers.WechatAppQrCodeSerializer
    permission_classes = [IsAuthenticated]

    def create(self, request, *args, **kwargs):
        from system.models import WechatConfig
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = request.data
        sys_id = data.get('sys_id')
        if sys_id is None:
            return Response({
                "error": "sys_id 不能为空"
            }, status=400)
        weconf = WechatConfig.objects.get(system__sys_id=data['sys_id'], wxa_aid__isnull=False)
        appid = getattr(weconf, 'wxa_aid')
        secret = getattr(weconf, 'wxa_sk')
        if not appid or not secret:
            return Response({
                "error": "小程序未配置"
            }, status=400)
        redis_client = get_redis_connection("default")
        session_interface = RedisStorage(
            redis_client,
            prefix=f"wechatpy_{sys_id}",
        )
        cli = wechatpy.WeChatClient(
            appid,
            secret,
            session=session_interface
        )
        try:
            qrcode = cli.wxa.get_wxa_code_unlimited(
                scene=data['scene'], page=data.get('page', 'pages/index/index')
            )
        except wechatpy.exceptions.WeChatClientException as e:
            return Response({
                "error": str(e)
            }, status=500)
        # save qrcode image content to django file
        date = datetime.datetime.now()
        file_name = f'wxaqrcode/{date.year}/{date.month}/wxacode_{sys_id}_{int(time.time())}.png'
        file_path = default_storage.save(file_name, io.BytesIO(qrcode.content))
        file_url = default_storage.url(file_path)

        return Response({
            "qrcode": file_url,
        })


class WechatAppPhoneNum(GenericViewSet):
    """
    微信小程序获取手机号接口

    用于获取微信小程序用户绑定的手机号码信息

    请求参数:
        sys_id (int): 系统ID
        code (str): 手机号获取凭证

    返回数据:
        微信接口返回的手机号信息
    """
    queryset = BaseConfigFileUpload.objects.none()
    serializer_class = serializers.WechatAppPhoneNumSerializer
    permission_classes = [IsAuthenticated]

    def create(self, request, *args, **kwargs):
        from system.models import WechatConfig
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = request.data
        sys_id = data.get('sys_id')
        if sys_id is None:
            return Response({
                "error": "sys_id 不能为空"
            }, status=400)
        code = data.get('code')
        if code is None:
            return Response({
                "error": "code 不能为空"
            }, status=400)
        weconf = WechatConfig.objects.get(system__sys_id=data['sys_id'], wxa_aid__isnull=False)
        appid = getattr(weconf, 'wxa_aid')
        secret = getattr(weconf, 'wxa_sk')
        if not appid or not secret:
            return Response({
                "error": "小程序未配置"
            }, status=400)
        redis_client = get_redis_connection("default")
        session_interface = RedisStorage(
            redis_client,
            prefix=f"wechatpy_{sys_id}",
        )
        cli = wechatpy.WeChatClient(
            appid,
            secret,
            session=session_interface
        )
        try:
            phone = cli.wxa._post(
                'wxa/business/getuserphonenumber',
                data={
                    'code': code,
                }
            )
        except wechatpy.exceptions.WeChatClientException as e:
            return Response({
                "error": str(e)
            }, status=500)
        return Response(phone)


class PinYinViewSet(GenericViewSet):
    """
    汉字转拼音接口

    将中文字符串转换为对应的拼音数组

    请求参数:
        origin_string (str): 待转换的中文字符串,最大长度64

    返回数据:
        {
            "pinyin": ["pin", "yin", "shu", "zu"]
        }
    """
    queryset = BaseConfigFileUpload.objects.none()
    serializer_class = serializers.PinYinSerializer
    permission_classes = [IsAuthenticatedOrReadOnly]

    def create(self, request, *args, **kwargs):
        """字符串转拼音数组API"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.save()
        return Response(data)



def get_token_plus_url(originurl, token):
    pr = urllib.parse.urlparse(originurl)
    fragment = pr.fragment
    ppr = urllib.parse.urlparse(fragment)
    if '?' in ppr.query:
        query = ppr.query + '&token=' + token
    else:
        query = '?token=' + token
    ppr_list = list(ppr)
    ppr_list[4] = query
    fragment_res = urllib.parse.urlunparse(ppr_list)
    pr_list = list(pr)
    pr_list[5] = fragment_res
    return urllib.parse.urlunparse(pr_list)


class MPTemplateMsgSendAPI(GenericViewSet):
    """
    微信公众号模板消息发送接口

    发送微信公众号模板消息

    请求参数:
        sys_id (int): 系统ID
        appid (str, optional): 公众号APPID
        touser (str): 接收人openid
        template_id (str): 模板消息模板ID
        url (str, optional): 模板消息跳转URL
        data (dict): 模板消息内容字段数据

    返回数据:
        {
            "errcode": "微信错误码,正常返回0",
            "errmsg": "微信错误信息,正常返回ok"
        }
    """
    queryset = BaseConfigFileUpload.objects.none()
    serializer_class = serializers.MPTemplateMsgSerializer
    permission_classes = [AllowAny]

    def create(self, request, *args, **kwargs):
        """
        发送微信公众号模板消息API，
        需要首先配置相应系统的 SystemWechatConfig

        正常时的返回 JSON 数据包示例：
        ```
        {
           "errcode" : 0,
           "errmsg" : "ok"
        }
        ```
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=201)


class WXATemplateMsgSendAPI(GenericViewSet):
    """
    微信小程序模板消息发送接口

    发送微信小程序订阅消息

    请求参数:
        sys_id (int): 系统ID
        appid (str, optional): 小程序APPID
        touser (str): 接收人openid
        template_id (str): 模板消息模板ID
        page (str, optional): 消息跳转page
        data (dict): 模板消息内容字段数据

    返回数据:
        {
            "errcode": "微信错误码,正常返回0",
            "errmsg": "微信错误信息,正常返回ok"
        }
    """
    queryset = BaseConfigFileUpload.objects.none()
    serializer_class = serializers.WXATemplateMsgSerializer
    permission_classes = [AllowAny]

    def create(self, request, *args, **kwargs):
        """
        发送微信小程序模板消息API，
        需要首先配置相应系统的 SystemWechatConfig

        正常时的返回 JSON 数据包示例：
        ```
        {
           "errcode" : 0,
           "errmsg" : "ok"
        }
        ```
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=201)


class SMSSendCodeAPI(GenericViewSet):
    """
    短信验证码发送接口

    发送手机短信验证码

    请求参数:
        sys_id (int): 系统ID
        org_id (int, optional): 机构ID
        phone (str): 手机号码
        expire_seconds (int, optional): 超时时长

    返回数据:
        {
            "errcode": "错误码,0表示成功",
            "errmsg": "错误信息"
        }
    """
    queryset = SMSLog.objects.none()
    serializer_class = serializers.SMSSendCodeSerializer
    permission_classes = [IsAuthenticated]

    def create(self, request, *args, **kwargs):
        """
        发送短信验证码API，

        正常时的返回 JSON 数据包示例：
        ```
        {
           "errcode" : 0,
           "errmsg" : "ok"
        }
        ```
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=201)


class SMSCodeValidAPI(GenericViewSet):
    """
    短信验证码校验接口

    验证用户输入的短信验证码是否正确

    请求参数:
        sys_id (int): 系统ID
        org_id (int, optional): 机构ID
        phone (str): 手机号码
        code (str): 手机验证码

    返回数据:
        {
            "errcode": "错误码,0表示成功",
            "errmsg": "错误信息"
        }
    """
    queryset = SMSLog.objects.none()
    serializer_class = serializers.SMSCodeValidSerializer
    permission_classes = [IsAuthenticated]

    def create(self, request, *args, **kwargs):
        """
        短信验证码校验API，
        需要首先配置相应系统的 SMSConfig

        正常时的返回 JSON 数据包示例：
        ```
        {
           "errcode" : 0,
           "errmsg" : "ok"
        }
        ```
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=201)


@csrf_exempt
def exportsvg(request):
    import re
    from zipfile import ZipFile
    title = request.POST.get('title', 'pictures')
    images = {}
    sub_titles = {}

    # 添加图片
    for key in request.POST.keys():
        if key.startswith("img_"):
            images[int(key[4:])] = request.POST[key]
        if key.startswith("title_"):
            sub_titles[int(key[6:])] = request.POST[key]
    response = HttpResponse(content_type='application/zip')
    zf = ZipFile(response, 'w')
    numbs = list(images.keys())
    numbs.sort()
    for num in numbs:
        sub_title = sub_titles.get(num)
        if not sub_title:
            sub_title = str(num)
        img = images.get(num)
        if img:
            pat = re.compile(r'<\s*div[^>]*>[^<]*<\s*/\s*div\s*>', re.I)
            img = pat.sub('', img)
            zf.writestr(sub_title + '.svg', img)
    response['Content-Disposition'] = 'attachment;filename={}.zip'.format(title).encode('utf-8')
    return response


@csrf_exempt
def exportword(request):
    from io import BytesIO
    from docx import Document
    from docx.oxml.ns import qn
    from .utils import add_heading, add_zhengwen, add_pic
    title = request.POST.get('title', 'document')
    images = {}
    descriptions = {}
    sub_titles = {}

    doc = Document()
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 添加标题页
    add_heading(doc, title, 0, True)

    # 添加图片
    for key in request.POST.keys():
        if key.startswith("img_"):
            images[int(key[4:])] = request.POST[key]
        if key.startswith("desc_"):
            descriptions[int(key[5:])] = request.POST[key]
        if key.startswith("title_"):
            sub_titles[int(key[6:])] = request.POST[key]

    numbs = list(images.keys())
    numbs.sort()
    for num in numbs:
        sub_title = sub_titles.get(num)
        if sub_title:
            add_heading(doc, sub_title, 2)
        img = images.get(num)
        if img:
            add_pic(doc, img)
        desc = descriptions.get(num)
        if desc:
            add_zhengwen(doc, desc)

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    response = HttpResponse(content_type='application/msword')
    file_name = title
    response['Content-Disposition'] = 'attachment;filename={}.docx'.format(file_name).encode('utf-8')
    response.write(stream.getvalue())
    return response


@csrf_exempt
def docx_to_pdf(request):
    if request.method != 'POST':
        return HttpResponse('', status=400)
    file_content = request.FILES['file']
    upload_to = 'document/%Y/%m/%d/%H-%M-%S/'
    upload_to = datetime.datetime.now().strftime(upload_to)
    path = default_storage.save(upload_to + file_content.name, file_content)
    abs_path = default_storage.path(path)
    pdf_file = office.convert_to(os.path.dirname(abs_path), abs_path)
    f_name = os.path.split(pdf_file)[1]
    return HttpResponse(json.dumps({
        'pdf': f'{upload_to}{f_name}',
        'docx': f'{upload_to}{file_content.name}'
    }))


@csrf_exempt
def save_sign_pdf(request):
    if request.method != 'POST':
        return HttpResponse('', status=400)
    print(request.FILES)
    file_content = request.FILES['file']
    origin_path = request.POST['origin'].replace(' ', '+')
    print(origin_path)
    origin_path = base64.b64decode(origin_path).decode('utf-8')
    file_name = 'document/' + origin_path.split('document/')[1]
    # upload_to = 'document/%Y/%m/%d/%H-%M-%S/'
    # upload_to = datetime.datetime.now().strftime(upload_to)
    if default_storage.exists(file_name):
        default_storage.delete(file_name)
        print(f'delete {file_name}')
    path = default_storage.save(file_name, file_content)
    abs_path = default_storage.path(path)
    print(abs_path)
    return HttpResponse('ok')


@csrf_exempt
def add_image_to_pdf_api(request):
    """
    给PDF文件添加图片API

    参数：
    {
     "file_id": "",
     "page_number": 2,
     "x": 120,
     "y": 240,
     "width": 300,
     "height": 120,
     "image_base64": "..."
    }
    返回值：
    {
        "success": true
    }
    """
    from django.http import JsonResponse
    from baseconfig.models import BaseConfigFileUpload
    from .pdfutils import add_image_to_pdf
    if request.method != 'POST':
        return HttpResponse('', status=400)
    try:
        json_data = json.loads(request.body.decode())
        pdf_file_id = json_data.get('file_id')
        page_number = int(json_data.get('page_number'))
        x = int(json_data.get('x'))
        y = int(json_data.get('y'))
        width = int(json_data.get('width'))
        height = int(json_data.get('height'))
        image_base64 = json_data.get('image_base64')
    except (json.JSONDecodeError, TypeError) as e:
        return JsonResponse({'success': False, 'message': f'JSON数据解析失败: {e}'}, status=400)
    try:
        pdf_path = BaseConfigFileUpload.objects.get(id=pdf_file_id).file.path
    except BaseConfigFileUpload.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'PDF文件不存在'}, status=404)
    output_path = pdf_path

    add_image_to_pdf(
        pdf_path, page_number, x, y, width, height, image_base64, output_path,
        use_tmp_files=True
    )

    return JsonResponse({'success': True})
