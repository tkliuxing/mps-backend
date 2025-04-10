import glob
import os
import re
from io import BytesIO

from reportlab.lib import colors
from reportlab.pdfbase.pdfmetrics import registerFont
from svglib.svglib import svg2rlg
import matplotlib.font_manager
from matplotlib import pyplot as plt
from reportlab.graphics import renderPM
from reportlab.pdfbase.ttfonts import TTFont
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


if os.path.exists('''/usr/share/fonts/chinese/'''):
    ttfs = glob.glob('/usr/share/fonts/chinese/*.ttf')
    for ttf in ttfs:
        matplotlib.font_manager.fontManager.addfont(ttf)


plt.rcParams['font.sans-serif'] = ['Microsoft YaHei']
# plt.rcParams['font.family'] = ['WenQuanYi Micro Hei Mono']
plt.rcParams['axes.unicode_minus'] = False

if os.path.exists('/usr/share/fonts/msyh.ttf'):
    matplotlib.font_manager.fontManager.addfont('/usr/share/fonts/msyh.ttf')
    registerFont(TTFont('Microsoft YaHei', '/usr/share/fonts/msyh.ttf'))

def add_heading(document, text, level, center=False):
    head = document.add_heading(' ', level)
    if center:
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = head.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_zhengwen(document, text, nopadding=False):
    paragraph = document.add_paragraph(text)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Inches(0.3)
    if nopadding:
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)


def svg_to_png(svg_code):
    bio = BytesIO()
    bio.write(svg_code.encode('utf-8'))
    bio.seek(0)
    draw = svg2rlg(bio)
    bio.close()
    del bio
    fio = BytesIO()
    renderPM.drawToFile(draw, fn=fio, dpi=72, fmt="PNG", bg=colors.HexColor('#FFFFFF'))
    fio.seek(0)
    return fio


def add_pic(document, img_str):
    img_str = img_str.replace(
        "fill: transparent;", "fill: #FFFFFF;"
    ).replace(
        "sans-serif;", ""
    ).replace(
        'style="user-select: none;',
        'style="user-select: none;font-family:\'Microsoft YaHei\';'
    )
    pat = re.compile('<\s*div[^>]*>[^<]*<\s*/\s*div\s*>', re.I)
    img_str = pat.sub('', img_str)
    pic = svg_to_png(img_str)
    document.add_picture(pic, width=Inches(5.9))
